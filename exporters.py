"""
Export module — generates PDF, summary PDF, Word (.docx), and CSV reports
from an AnalysisResult without any external API or AI.
"""

import io
import csv
from datetime import datetime
from typing import Optional

from analyzer import AnalysisResult


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

RISK_COLOR = {
    "Low":    (76,  175, 132),   # green
    "Medium": (244, 200,  66),   # yellow
    "High":   (255, 107, 122),   # red
}

GOLD    = (212, 175,  55)
DARK    = ( 13,  13,  13)
WHITE   = (240, 235, 225)
GREY    = (100, 100, 100)
LGREY   = (220, 220, 220)

def _now() -> str:
    return datetime.now().strftime("%B %d, %Y at %H:%M")

def _risk_icon(level: str) -> str:
    return {"Low": "✓", "Medium": "!", "High": "✕"}.get(level, "?")


# ─────────────────────────────────────────────────────────────────────────────
# Full PDF report  (ReportLab)
# ─────────────────────────────────────────────────────────────────────────────

def export_pdf(result: AnalysisResult) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=18*mm, bottomMargin=18*mm,
        title="T&C Analysis Report"
    )

    W, H = A4
    cw = W - 40*mm  # content width

    # ── Colour helpers ──────────────────────────────────────────────────────
    def rgb(t):  return colors.Color(*[v/255 for v in t])

    rc  = rgb(RISK_COLOR.get(result.risk_level, GREY))
    gold_c = rgb(GOLD)
    dark_c = rgb(DARK)
    grey_c = rgb(GREY)
    lgrey_c = rgb(LGREY)

    # ── Styles ──────────────────────────────────────────────────────────────
    base = getSampleStyleSheet()

    def sty(name, parent="Normal", **kw):
        s = ParagraphStyle(name, parent=base[parent], **kw)
        return s

    s_title   = sty("title",   fontSize=22, leading=28, textColor=dark_c, spaceAfter=4, fontName="Helvetica-Bold")
    s_sub     = sty("sub",     fontSize=11, leading=16, textColor=grey_c,  spaceAfter=12)
    s_h2      = sty("h2",      fontSize=13, leading=18, textColor=dark_c,  spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold")
    s_body    = sty("body",    fontSize=9,  leading=14, textColor=dark_c,  spaceAfter=4)
    s_small   = sty("small",   fontSize=8,  leading=12, textColor=grey_c,  spaceAfter=2)
    s_badge   = sty("badge",   fontSize=8,  leading=12, textColor=grey_c)
    s_ev      = sty("ev",      fontSize=8,  leading=12, textColor=grey_c,  leftIndent=12, spaceAfter=3)

    story = []

    # ── Header ──────────────────────────────────────────────────────────────
    header_data = [[
        Paragraph("⚖ Terms &amp; Conditions Analysis Report", s_title),
        Paragraph(f"Generated {_now()}", s_small),
    ]]
    header_tbl = Table(header_data, colWidths=[cw*0.75, cw*0.25])
    header_tbl.setStyle(TableStyle([
        ("VALIGN",    (0,0), (-1,-1), "BOTTOM"),
        ("ALIGN",     (1,0), (1,0),   "RIGHT"),
        ("BACKGROUND",(0,0), (-1,-1), colors.white),
        ("BOX",       (0,0), (-1,-1), 0, colors.white),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(header_tbl)
    story.append(HRFlowable(width="100%", thickness=2, color=gold_c, spaceAfter=12))

    # ── Doc type & summary ──────────────────────────────────────────────────
    story.append(Paragraph(result.document_type.upper(), s_badge))
    story.append(Paragraph(result.document_summary, s_body))
    story.append(Spacer(1, 6))

    # ── Risk banner ─────────────────────────────────────────────────────────
    risk_data = [[
        Paragraph(f"<b>{_risk_icon(result.risk_level)}  {result.risk_level} Risk</b>", sty("rk", fontSize=14, textColor=rc, fontName="Helvetica-Bold")),
        Paragraph(result.risk_reason, sty("rr", fontSize=9, leading=13, textColor=dark_c)),
        Paragraph(f"<b>{result.risk_score}/100</b>", sty("rs", fontSize=14, textColor=rc, fontName="Helvetica-Bold", alignment=2)),
    ]]
    risk_tbl = Table(risk_data, colWidths=[cw*0.2, cw*0.6, cw*0.2])
    risk_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), rgb(tuple(int(v*0.15) for v in RISK_COLOR.get(result.risk_level, GREY)))),
        ("BOX",        (0,0), (-1,-1), 1.5, rc),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING",(0,0), (-1,-1), 10),
        ("RIGHTPADDING",(0,0), (-1,-1), 10),
        ("TOPPADDING", (0,0), (-1,-1), 10),
        ("BOTTOMPADDING",(0,0), (-1,-1), 10),
        ("ROUNDEDCORNERS", [6]),
    ]))
    story.append(KeepTogether([risk_tbl]))
    story.append(Spacer(1, 14))

    # ── Key Points ──────────────────────────────────────────────────────────
    story.append(Paragraph("Key Points to Know", s_h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c, spaceAfter=8))

    for kp in result.key_points:
        bg = rgb((255, 193, 7, 15)) if kp.watch_out else rgb((245, 245, 245))
        border_c = rgb((255, 193, 7)) if kp.watch_out else lgrey_c

        rows = [[
            Paragraph(kp.icon, sty("ico", fontSize=16, leading=20)),
            [
                Paragraph(f"<font color='#888' size='7'>{kp.category.upper()}</font>", s_small),
                Paragraph(f"<b>{kp.title}</b>{'  <font color=\"#f4c842\" size=\"7\">⚠ WATCH OUT</font>' if kp.watch_out else ''}", sty("kt", fontSize=9, leading=13, fontName="Helvetica-Bold")),
                Paragraph(kp.detail, s_body),
            ] + ([Paragraph(f"<i>&ldquo;{ev}&rdquo;</i>", s_ev) for ev in kp.evidence[:1]] if kp.evidence else []),
        ]]
        tbl = Table(rows, colWidths=[14*mm, cw - 14*mm])
        tbl.setStyle(TableStyle([
            ("VALIGN",      (0,0), (-1,-1), "TOP"),
            ("BACKGROUND",  (0,0), (-1,-1), colors.white),
            ("BOX",         (0,0), (-1,-1), 0.75, lgrey_c),
            ("LEFTPADDING", (0,0), (-1,-1), 8),
            ("RIGHTPADDING",(0,0), (-1,-1), 8),
            ("TOPPADDING",  (0,0), (-1,-1), 8),
            ("BOTTOMPADDING",(0,0), (-1,-1), 8),
            ("ROUNDEDCORNERS", [4]),
        ]))
        story.append(KeepTogether([tbl, Spacer(1, 5)]))

    story.append(Spacer(1, 6))

    # ── Before Signing ──────────────────────────────────────────────────────
    story.append(Paragraph("Before You Sign", s_h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c, spaceAfter=8))

    for i, item in enumerate(result.before_signing, 1):
        row = [[
            Paragraph(f"<b>{i}</b>", sty("cn", fontSize=9, textColor=gold_c, fontName="Helvetica-Bold", alignment=1)),
            Paragraph(item, s_body),
        ]]
        t = Table(row, colWidths=[10*mm, cw - 10*mm])
        t.setStyle(TableStyle([
            ("VALIGN",       (0,0), (-1,-1), "TOP"),
            ("LEFTPADDING",  (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ("TOPPADDING",   (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0), (-1,-1), 4),
            ("LINEBELOW",    (0,0), (-1,0),  0.3, lgrey_c),
        ]))
        story.append(t)

    story.append(Spacer(1, 10))

    # ── Red Flags ───────────────────────────────────────────────────────────
    story.append(Paragraph("Red Flags", s_h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c, spaceAfter=8))

    if result.red_flags:
        for rf in result.red_flags:
            items = [
                [Paragraph("🚩", sty("fi", fontSize=10, leading=14)),
                 Paragraph(rf.message, s_body)]
            ]
            if rf.evidence:
                items.append([
                    Paragraph("", s_small),
                    Paragraph(f"<i>&ldquo;{rf.evidence[0][:200]}&rdquo;</i>", s_ev),
                ])
            t = Table(items, colWidths=[10*mm, cw - 10*mm])
            t.setStyle(TableStyle([
                ("VALIGN",       (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING",  (0,0), (-1,-1), 6),
                ("TOPPADDING",   (0,0), (-1,-1), 3),
                ("BOTTOMPADDING",(0,0), (-1,-1), 3),
                ("LINEBELOW",    (0,0), (-1,-1), 0.3, rgb((220, 53, 69))),
            ]))
            story.append(t)
    else:
        story.append(Paragraph("No major red flags detected.", s_small))

    story.append(Spacer(1, 10))

    # ── Readability ──────────────────────────────────────────────────────────
    if result.readability:
        rd = result.readability
        story.append(Paragraph("Readability Analysis", s_h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c, spaceAfter=8))
        story.append(Paragraph(f"<b>{rd.grade_label}</b> — {rd.ease_label}", s_body))
        story.append(Spacer(1, 4))

        metrics = [
            ["Metric", "Value", "What it means"],
            ["Flesch Reading Ease", f"{rd.flesch_ease}/100", "Higher = easier (80+ is plain English)"],
            ["Flesch-Kincaid Grade", f"Grade {rd.flesch_grade}", "US school grade level required"],
            ["Gunning Fog Index",   f"{rd.gunning_fog}",       "Years of education needed to understand"],
            ["Avg Sentence Length", f"{rd.avg_sentence_len} words", "Legal docs often run 20–40 words/sentence"],
            ["Complex Words",       f"{rd.complex_word_pct}%",     "% of words with 3+ syllables"],
        ]
        mt = Table(metrics, colWidths=[cw*0.3, cw*0.2, cw*0.5])
        mt.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  rgb(DARK)),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
            ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,-1), 8),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, rgb(LGREY)]),
            ("GRID",        (0,0), (-1,-1), 0.3, lgrey_c),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING",(0,0), (-1,-1), 6),
            ("TOPPADDING",  (0,0), (-1,-1), 5),
            ("BOTTOMPADDING",(0,0), (-1,-1), 5),
        ]))
        story.append(mt)

    # ── Footer ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c))
    story.append(Paragraph(
        "⚠ This report is for informational purposes only and does not constitute legal advice. "
        "For important agreements, consult a qualified legal professional.",
        sty("foot", fontSize=7, leading=10, textColor=grey_c, spaceAfter=0)
    ))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Summary PDF  — clean one-pager
# ─────────────────────────────────────────────────────────────────────────────

def export_summary_pdf(result: AnalysisResult) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=22*mm, rightMargin=22*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )
    W, _ = A4
    cw = W - 44*mm

    def rgb(t): return colors.Color(*[v/255 for v in t])

    rc   = rgb(RISK_COLOR.get(result.risk_level, GREY))
    gold_c = rgb(GOLD)
    lgrey_c = rgb(LGREY)
    grey_c  = rgb(GREY)

    def sty(**kw): return ParagraphStyle("s", **kw)

    story = []

    # ── Title ────────────────────────────────────────────────────────────────
    story.append(Paragraph(
        f"<b>T&amp;C Summary</b> &nbsp;<font color='#888' size='9'>— {result.document_type}</font>",
        sty(fontSize=18, leading=24, fontName="Helvetica-Bold", spaceAfter=2)))
    story.append(Paragraph(f"<font size='8' color='#888'>Generated {_now()}</font>",
        sty(fontSize=8, leading=12, spaceAfter=8)))
    story.append(HRFlowable(width="100%", thickness=2, color=gold_c, spaceAfter=10))

    # ── Summary ──────────────────────────────────────────────────────────────
    story.append(Paragraph(result.document_summary,
        sty(fontSize=9, leading=14, spaceAfter=10)))

    # ── Risk pill ────────────────────────────────────────────────────────────
    risk_row = [[
        Paragraph(f"<b>{result.risk_level} Risk — {result.risk_score}/100</b>",
            sty(fontSize=13, fontName="Helvetica-Bold", textColor=rc, leading=18)),
        Paragraph(result.risk_reason,
            sty(fontSize=8, leading=13, spaceAfter=0)),
    ]]
    rt = Table(risk_row, colWidths=[cw*0.32, cw*0.68])
    rt.setStyle(TableStyle([
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("BACKGROUND",  (0,0), (-1,-1), rgb((240,240,240))),
        ("BOX",         (0,0), (-1,-1), 1.5, rc),
        ("LEFTPADDING", (0,0), (-1,-1), 10),
        ("TOPPADDING",  (0,0), (-1,-1), 8),
        ("BOTTOMPADDING",(0,0), (-1,-1), 8),
    ]))
    story.append(rt)
    story.append(Spacer(1, 12))

    # ── Top key points (first 5) ─────────────────────────────────────────────
    story.append(Paragraph("<b>Key Points</b>",
        sty(fontSize=11, fontName="Helvetica-Bold", spaceAfter=6)))

    kp_data = [["", "Topic", "Summary"]]
    for kp in result.key_points[:5]:
        kp_data.append([
            Paragraph(kp.icon, sty(fontSize=11, leading=14)),
            Paragraph(f"<b>{kp.title}</b>{'  ⚠' if kp.watch_out else ''}",
                sty(fontSize=8, leading=12, fontName="Helvetica-Bold")),
            Paragraph(kp.detail, sty(fontSize=8, leading=12)),
        ])

    kp_tbl = Table(kp_data, colWidths=[8*mm, cw*0.25, cw*0.65])
    kp_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0),  rgb(DARK)),
        ("TEXTCOLOR",    (0,0), (-1,0),  colors.white),
        ("FONTNAME",     (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 8),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.white, rgb(LGREY)]),
        ("GRID",         (0,0), (-1,-1), 0.3, lgrey_c),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING",  (0,0), (-1,-1), 6),
        ("TOPPADDING",   (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0), (-1,-1), 4),
    ]))
    story.append(kp_tbl)
    story.append(Spacer(1, 12))

    # ── Top red flags ─────────────────────────────────────────────────────────
    story.append(Paragraph("<b>Red Flags</b>",
        sty(fontSize=11, fontName="Helvetica-Bold", spaceAfter=6)))

    if result.red_flags:
        for rf in result.red_flags[:4]:
            story.append(Paragraph(f"🚩  {rf.message}",
                sty(fontSize=8, leading=13, spaceAfter=3)))
    else:
        story.append(Paragraph("No major red flags detected.",
            sty(fontSize=8, leading=12, textColor=grey_c)))

    story.append(Spacer(1, 12))

    # ── Checklist (top 3) ────────────────────────────────────────────────────
    story.append(Paragraph("<b>Before You Sign</b>",
        sty(fontSize=11, fontName="Helvetica-Bold", spaceAfter=6)))

    for i, item in enumerate(result.before_signing[:3], 1):
        story.append(Paragraph(f"<b>{i}.</b>  {item}",
            sty(fontSize=8, leading=13, spaceAfter=4, leftIndent=6)))

    # ── Readability badge ────────────────────────────────────────────────────
    if result.readability:
        story.append(Spacer(1, 8))
        rd = result.readability
        story.append(Paragraph(
            f"<b>Readability:</b> {rd.grade_label} &nbsp;·&nbsp; "
            f"Flesch Ease {rd.flesch_ease}/100 &nbsp;·&nbsp; "
            f"Grade Level {rd.flesch_grade} &nbsp;·&nbsp; "
            f"Avg sentence {rd.avg_sentence_len} words",
            sty(fontSize=8, leading=13, textColor=grey_c)))

    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=0.5, color=lgrey_c))
    story.append(Paragraph(
        "This summary is for informational purposes only and does not constitute legal advice.",
        sty(fontSize=7, leading=10, textColor=grey_c)))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Word (.docx) export
# ─────────────────────────────────────────────────────────────────────────────

def export_word(result: AnalysisResult) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def add_heading(text, level=1, color=None):
        h = doc.add_heading(text, level=level)
        if color:
            for run in h.runs:
                run.font.color.rgb = RGBColor(*color)
        return h

    def add_para(text="", bold=False, italic=False, color=None, size=10, indent=0):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(text)
        run.bold, run.italic = bold, italic
        run.font.size = Pt(size)
        if color: run.font.color.rgb = RGBColor(*color)
        return p

    def add_rule():
        doc.add_paragraph("─" * 80).runs[0].font.color.rgb = RGBColor(*LGREY)

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Terms & Conditions Analysis Report", 0)
    title.runs[0].font.color.rgb = RGBColor(*DARK)

    add_para(f"Generated: {_now()}", color=GREY, size=9)
    add_para(f"Document type: {result.document_type}", bold=True)
    add_para(result.document_summary, size=9)
    add_rule()

    # ── Risk ─────────────────────────────────────────────────────────────────
    rc = RISK_COLOR.get(result.risk_level, GREY)
    add_heading("Risk Assessment", 1)
    p = doc.add_paragraph()
    run = p.add_run(f"{result.risk_level} Risk  ({result.risk_score}/100)")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = RGBColor(*rc)
    add_para(result.risk_reason, size=9)

    # ── Key Points ───────────────────────────────────────────────────────────
    add_heading("Key Points to Know", 1)
    for kp in result.key_points:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{kp.icon}  {kp.title}")
        run.bold = True; run.font.size = Pt(10)
        if kp.watch_out:
            w = p.add_run("  ⚠ Watch Out")
            w.font.color.rgb = RGBColor(244, 200, 66); w.font.size = Pt(8)
        cat_p = doc.add_paragraph()
        cat_p.paragraph_format.left_indent = Inches(0.25)
        cat_run = cat_p.add_run(f"{kp.category}  ·  {kp.detail}")
        cat_run.font.size = Pt(9)
        if kp.evidence:
            ev_p = doc.add_paragraph()
            ev_p.paragraph_format.left_indent = Inches(0.25)
            ev_run = ev_p.add_run(f'"{kp.evidence[0][:200]}"')
            ev_run.italic = True; ev_run.font.size = Pt(8)
            ev_run.font.color.rgb = RGBColor(*GREY)

    # ── Before Signing ───────────────────────────────────────────────────────
    add_heading("Before You Sign", 1)
    for i, item in enumerate(result.before_signing, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item).font.size = Pt(9)

    # ── Red Flags ────────────────────────────────────────────────────────────
    add_heading("Red Flags", 1)
    if result.red_flags:
        for rf in result.red_flags:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"🚩  {rf.message}")
            run.font.size = Pt(9); run.font.color.rgb = RGBColor(220, 53, 69)
            if rf.evidence:
                ev_p = doc.add_paragraph()
                ev_p.paragraph_format.left_indent = Inches(0.25)
                ev_run = ev_p.add_run(f'"{rf.evidence[0][:200]}"')
                ev_run.italic = True; ev_run.font.size = Pt(8)
                ev_run.font.color.rgb = RGBColor(*GREY)
    else:
        add_para("No major red flags detected.", color=GREY, size=9)

    # ── Readability ──────────────────────────────────────────────────────────
    if result.readability:
        rd = result.readability
        add_heading("Readability Analysis", 1)
        add_para(f"{rd.grade_label} — {rd.ease_label}", bold=True)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Value", "Notes"
        metrics = [
            ("Flesch Reading Ease",  f"{rd.flesch_ease}/100",  "80+ is plain English"),
            ("Flesch-Kincaid Grade", f"Grade {rd.flesch_grade}", "US school grade level"),
            ("Gunning Fog Index",    str(rd.gunning_fog),       "Years of education needed"),
            ("Avg Sentence Length",  f"{rd.avg_sentence_len} words", "Legal avg is often 20–40"),
            ("Complex Word %",       f"{rd.complex_word_pct}%", "Words with 3+ syllables"),
        ]
        for m, v, n in metrics:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text = m, v, n

    # ── Disclaimer ───────────────────────────────────────────────────────────
    add_rule()
    add_para("⚠ This report is for informational purposes only and does not constitute legal advice.",
             italic=True, color=GREY, size=8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# CSV export
# ─────────────────────────────────────────────────────────────────────────────

def export_csv(result: AnalysisResult) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf)

    # ── Summary ──────────────────────────────────────────────────────────────
    w.writerow(["SECTION", "FIELD", "VALUE"])
    w.writerow(["Summary", "Document Type",   result.document_type])
    w.writerow(["Summary", "Risk Level",      result.risk_level])
    w.writerow(["Summary", "Risk Score",      result.risk_score])
    w.writerow(["Summary", "Risk Reason",     result.risk_reason])
    w.writerow(["Summary", "Word Count",      result.word_count])
    w.writerow(["Summary", "Char Count",      result.char_count])
    w.writerow(["Summary", "Summary",         result.document_summary])
    w.writerow([])

    # ── Readability ──────────────────────────────────────────────────────────
    if result.readability:
        rd = result.readability
        w.writerow(["Readability", "Grade Label",       rd.grade_label])
        w.writerow(["Readability", "Flesch Ease",       rd.flesch_ease])
        w.writerow(["Readability", "Flesch Grade",      rd.flesch_grade])
        w.writerow(["Readability", "Gunning Fog",       rd.gunning_fog])
        w.writerow(["Readability", "Avg Sentence Len",  rd.avg_sentence_len])
        w.writerow(["Readability", "Avg Word Len",      rd.avg_word_len])
        w.writerow(["Readability", "Complex Word %",    rd.complex_word_pct])
    w.writerow([])

    # ── Key Points ───────────────────────────────────────────────────────────
    w.writerow(["KEY POINTS"])
    w.writerow(["Category", "Icon", "Title", "Detail", "Watch Out", "Evidence"])
    for kp in result.key_points:
        w.writerow([kp.category, kp.icon, kp.title, kp.detail,
                    "YES" if kp.watch_out else "NO",
                    " | ".join(kp.evidence)])
    w.writerow([])

    # ── Red Flags ─────────────────────────────────────────────────────────────
    w.writerow(["RED FLAGS"])
    w.writerow(["Message", "Evidence"])
    for rf in result.red_flags:
        w.writerow([rf.message, " | ".join(rf.evidence)])
    w.writerow([])

    # ── Checklist ─────────────────────────────────────────────────────────────
    w.writerow(["BEFORE SIGNING CHECKLIST"])
    w.writerow(["#", "Action"])
    for i, item in enumerate(result.before_signing, 1):
        w.writerow([i, item])

    return buf.getvalue().encode("utf-8-sig")  # BOM for Excel compatibility
